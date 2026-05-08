"""
Quiz Builder routes.

GET    /v1/quiz/sessions                                              list teacher's sessions
POST   /v1/quiz/sessions                                              create session
GET    /v1/quiz/sessions/{session_id}                                 get single session (polled during generation)
DELETE /v1/quiz/sessions/{session_id}                                 delete session (204)
POST   /v1/quiz/sessions/{session_id}/source                          save source material
POST   /v1/quiz/sessions/{session_id}/blueprint                       generate question blueprint via LLM
PATCH  /v1/quiz/sessions/{session_id}                                 update session metadata
PUT    /v1/quiz/sessions/{session_id}/blueprint                       save blueprint directly
POST   /v1/quiz/sessions/{session_id}/generate                        enqueue question generation job (202)
GET    /v1/quiz/sessions/{session_id}/questions                       list questions
POST   /v1/quiz/sessions/{session_id}/questions/reorder               reorder questions
PATCH  /v1/quiz/sessions/{session_id}/questions/{question_id}         update question
DELETE /v1/quiz/sessions/{session_id}/questions/{question_id}         delete question (204)
POST   /v1/quiz/sessions/{session_id}/questions/{question_id}/regenerate  regenerate via LLM
POST   /v1/quiz/sessions/{session_id}/questions/{question_id}/clone   clone question (201)
POST   /v1/quiz/sessions/{session_id}/questions/{question_id}/image   upload image
DELETE /v1/quiz/sessions/{session_id}/questions/{question_id}/image   remove image (204)
"""

from __future__ import annotations

import asyncio
from datetime import datetime, timezone
import html as _html
import json
import logging
import re
import uuid
from typing import Annotated, Any

from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
    Request,
    Response,
    UploadFile,
    status,
)
from pydantic import BaseModel

from gradenza_api.auth import AuthUser, require_roles
from gradenza_api.services.openrouter import call_openrouter, strip_json_fences
from gradenza_api.services.pdf_playwright import render_html_to_pdf_bytes
from gradenza_api.services.llm_structured import (
    StructuredJSONError,
    call_openrouter_structured_json,
)
from gradenza_api.services.quiz_schemas import (
    QuizQuestionSchema,
    validate_quiz_blueprint_items,
)
from gradenza_api.services.supabase_client import get_service_client
from gradenza_api.services.usage import AIUsage, record_ai_usage
from gradenza_api.settings import settings

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/v1/quiz", tags=["quiz"])

_TEACHER_ROLES = ("teacher", "tutor", "co_teacher", "school_admin")
_QUIZ_MODEL = "google/gemini-2.5-flash"

_DEFAULT_DIFFICULTY_DEFS: dict[str, str] = {
    "easy": "Straightforward recall or basic comprehension; suitable for checking core understanding.",
    "medium": "Requires explanation, comparison, or applying the content in a familiar context.",
    "challenge": "Requires deeper reasoning, synthesis, evaluation, or applying the content in a less familiar context.",
}

VALID_MODES = frozenset({"from_source", "from_prompt", "from_questions", "personalized"})
VALID_PRESETS = frozenset({
    "quick_check", "standard_worksheet", "exam_practice",
    "support_practice", "challenge_extension", "reading_comprehension",
})


# ── Helpers ────────────────────────────────────────────────────────────────────

def _fetch_session(session_id: str, teacher_id: str) -> dict | None:
    svc = get_service_client()
    result = (
        svc.table("quiz_builder_sessions")
        .select("*")
        .eq("id", session_id)
        .eq("teacher_id", teacher_id)
        .maybe_single()
        .execute()
    )
    return result.data


def _require_session(session_id: str, teacher_id: str) -> dict:
    session = _fetch_session(session_id, teacher_id)
    if not session:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Session not found")
    return session


_MC_ANSWER_RE = re.compile(r"^([A-Za-z])\.?$")


def _normalize_mc_answer(qj: dict[str, Any]) -> dict[str, Any]:
    """Uppercase-normalize correct_answer for MCQ questions stored via teacher edits.

    Mirrors the normalization in QuizQuestionSchema without the strict image_slot
    constraint, so it is safe to call on questions that already have an uploaded image.
    No-ops for non-MCQ questions or answers that don't look like a single letter.
    """
    if qj.get("question_type") != "multiple_choice":
        return qj
    raw = (qj.get("correct_answer") or "").strip()
    m = _MC_ANSWER_RE.match(raw)
    if m:
        return {**qj, "correct_answer": m.group(1).upper()}
    return qj


_OPTION_PREFIX_RE = re.compile(r"^\s*([A-Za-z])\.")


def _validate_mc_options(qj: dict[str, Any]) -> None:
    """Raise HTTPException(422) if MCQ options or correct_answer violate shape invariants.

    Allows questions that already have an uploaded image_slot (dict). Only validates
    when question_type is multiple_choice and options are explicitly provided.
    """
    if qj.get("question_type") != "multiple_choice":
        return
    options = qj.get("options")
    if options is None:
        return
    if not isinstance(options, list) or len(options) < 2:
        raise HTTPException(status_code=422, detail="MCQ questions must have at least 2 options")
    seen: set[str] = set()
    for opt in options:
        if not isinstance(opt, str):
            raise HTTPException(status_code=422, detail="Each option must be a string")
        m = _OPTION_PREFIX_RE.match(opt)
        if not m:
            raise HTTPException(
                status_code=422,
                detail=f"Option {opt!r} must start with a letter and dot, e.g. 'A. text'",
            )
        letter = m.group(1).upper()
        if letter in seen:
            raise HTTPException(status_code=422, detail=f"Duplicate option label {letter!r}")
        seen.add(letter)
    correct = (qj.get("correct_answer") or "").strip().upper().rstrip(".")
    if len(correct) != 1 or correct not in seen:
        raise HTTPException(
            status_code=422,
            detail=f"correct_answer must be a single letter matching one of the options (got {qj.get('correct_answer')!r})",
        )


# ── GET /v1/quiz/sessions ──────────────────────────────────────────────────────

@router.get("/sessions", summary="List teacher's quiz sessions")
async def list_sessions(
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> list[dict]:
    def _query() -> list[dict]:
        svc = get_service_client()
        result = (
            svc.table("quiz_builder_sessions")
            .select("*")
            .eq("teacher_id", user.id)
            .order("created_at", desc=True)
            .execute()
        )
        return result.data or []

    return await asyncio.to_thread(_query)


# ── POST /v1/quiz/sessions ─────────────────────────────────────────────────────

class CreateSessionRequest(BaseModel):
    mode: str
    preset: str
    subject: str | None = None
    grade_level: str | None = None
    title: str | None = None
    class_id: str | None = None


@router.post("/sessions", status_code=status.HTTP_201_CREATED, summary="Create a quiz session")
async def create_session(
    body: CreateSessionRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    if body.mode not in VALID_MODES:
        raise HTTPException(status_code=400, detail=f"Invalid mode: {body.mode!r}")
    if body.preset not in VALID_PRESETS:
        raise HTTPException(status_code=400, detail=f"Invalid preset: {body.preset!r}")

    if body.class_id is not None:
        def _check_class() -> bool:
            svc = get_service_client()
            accessible = _get_accessible_class_ids(svc, user.id, user.role)
            return body.class_id in accessible

        if not await asyncio.to_thread(_check_class):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="You do not have access to this class")

    def _insert() -> dict:
        svc = get_service_client()
        result = svc.table("quiz_builder_sessions").insert({
            "teacher_id": user.id,
            "mode": body.mode,
            "preset": body.preset,
            "subject": body.subject,
            "grade_level": body.grade_level,
            "title": body.title,
            "class_id": body.class_id,
            "status": "draft",
        }).execute()
        return result.data[0] if result.data else {}

    row = await asyncio.to_thread(_insert)
    if not row:
        raise HTTPException(status_code=500, detail="Failed to create session")
    return row


# ── GET /v1/quiz/sessions/{session_id} ────────────────────────────────────────

@router.get("/sessions/{session_id}", summary="Get a single quiz session")
async def get_session(
    session_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    return await asyncio.to_thread(_require_session, session_id, user.id)


# ── DELETE /v1/quiz/sessions/{session_id} ─────────────────────────────────────

@router.delete(
    "/sessions/{session_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete a quiz session",
)
async def delete_session(
    session_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> Response:
    await asyncio.to_thread(_require_session, session_id, user.id)

    def _delete() -> None:
        svc = get_service_client()
        svc.table("quiz_builder_sessions").delete().eq("id", session_id).execute()

    await asyncio.to_thread(_delete)
    return Response(status_code=status.HTTP_204_NO_CONTENT)


# ── POST /v1/quiz/sessions/{session_id}/source ────────────────────────────────

@router.post("/sessions/{session_id}/source", summary="Save source material")
async def save_source(
    session_id: str,
    source_type: Annotated[str, Form()],
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
    text: Annotated[str | None, Form()] = None,
    file: Annotated[UploadFile | None, File()] = None,
) -> dict:
    await asyncio.to_thread(_require_session, session_id, user.id)

    valid_types = {"pasted_text", "pdf", "txt", "docx", "prompt", "teacher_questions"}
    if source_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"Invalid source_type: {source_type!r}")

    raw_input: str = text or ""
    extracted_text: str = text or ""
    metadata: dict[str, Any] = {}

    if file is not None:
        file_bytes = await file.read()
        metadata["filename"] = file.filename
        metadata["size"] = len(file_bytes)

        if source_type == "txt":
            try:
                extracted_text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                extracted_text = file_bytes.decode("latin-1", errors="replace")
            raw_input = extracted_text
        elif source_type == "pdf":
            extracted_text = _extract_pdf_text(file_bytes)
            raw_input = extracted_text or f"[PDF: {file.filename}]"
        elif source_type == "docx":
            extracted_text = _extract_docx_text(file_bytes)
            raw_input = extracted_text or f"[DOCX: {file.filename}]"
        else:
            raw_input = f"[File: {file.filename}]"
            extracted_text = ""

    def _upsert() -> dict:
        svc = get_service_client()
        svc.table("quiz_source_materials").delete().eq("session_id", session_id).execute()
        result = svc.table("quiz_source_materials").insert({
            "session_id": session_id,
            "source_type": source_type,
            "raw_input": raw_input,
            "extracted_text": extracted_text,
            "metadata": metadata,
        }).execute()
        return result.data[0] if result.data else {}

    row = await asyncio.to_thread(_upsert)
    return {"ok": True, "id": row.get("id")}


# ── POST /v1/quiz/sessions/{session_id}/blueprint ─────────────────────────────

_BLUEPRINT_PROMPT = """\
You are an expert curriculum designer. Generate a quiz blueprint.

Session details:
- Mode: {mode}
- Preset: {preset}
- Subject: {subject}
- Grade level: {grade_level}
- Number of questions requested: {question_count}
- Question type split: {mcq_instruction}

Source material (first 2000 chars):
{source_text}

Preset guidelines:
- quick_check: mostly easy/medium, multiple_choice and fill_blank, 1 mark each
- standard_worksheet: mix of types, easy–medium, 1–3 marks each
- exam_practice: mix of all types, medium–challenge, 1–5 marks each
- support_practice: easy, short_answer and fill_blank, 1–2 marks each
- challenge_extension: medium–challenge, long_answer and short_answer, 2–5 marks each
- reading_comprehension: all short_answer, medium, 2–3 marks each

Return a JSON array of exactly {question_count} objects. Each:
{{"slot":<int 1..N>,"question_type":"multiple_choice"|"fill_blank"|"short_answer"|"long_answer","difficulty":"easy"|"medium"|"challenge","marks":<int 1..5>}}

Return ONLY the raw JSON array. No markdown, no explanation."""


class BlueprintRequest(BaseModel):
    question_count: int = 10
    mcq_count: int | None = None


@router.post("/sessions/{session_id}/blueprint", summary="Generate question blueprint via LLM")
async def generate_blueprint(
    session_id: str,
    body: BlueprintRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    session = await asyncio.to_thread(_require_session, session_id, user.id)

    def _get_source() -> dict | None:
        svc = get_service_client()
        result = (
            svc.table("quiz_source_materials")
            .select("extracted_text")
            .eq("session_id", session_id)
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        )
        return result.data[0] if result.data else None

    source = await asyncio.to_thread(_get_source)
    source_text = (source or {}).get("extracted_text") or "(no source text provided)"

    if body.mcq_count is not None:
        fr_count = body.question_count - body.mcq_count
        mcq_instruction = f"exactly {body.mcq_count} multiple_choice and {fr_count} free-response (fill_blank/short_answer/long_answer)"
    else:
        mcq_instruction = "choose an appropriate mix based on the preset"

    prompt = _BLUEPRINT_PROMPT.format(
        mode=session.get("mode", "from_prompt"),
        preset=session.get("preset", "standard_worksheet"),
        subject=session.get("subject") or "General",
        grade_level=session.get("grade_level") or "Secondary",
        question_count=body.question_count,
        mcq_instruction=mcq_instruction,
        source_text=source_text[:2000],
    )

    svc = get_service_client()
    _blueprint_route = f"POST /v1/quiz/sessions/{session_id}/blueprint"
    response_format = {
        "type": "json_schema",
        "json_schema": {
            "name": "quiz_blueprint",
            "strict": True,
            "schema": {
                "type": "array",
                "minItems": body.question_count,
                "maxItems": body.question_count,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "slot": {"type": "integer", "minimum": 1},
                        "question_type": {
                            "type": "string",
                            "enum": ["multiple_choice", "fill_blank", "short_answer", "long_answer"],
                        },
                        "difficulty": {"type": "string", "enum": ["easy", "medium", "challenge"]},
                        "marks": {"type": "integer", "minimum": 1, "maximum": 5},
                    },
                    "required": ["slot", "question_type", "difficulty", "marks"],
                },
            },
        },
    }

    base_messages = [
        {
            "role": "system",
            "content": "Return only valid JSON matching the requested schema. No markdown fences, no prose.",
        },
        {"role": "user", "content": prompt},
    ]
    retry_instruction = (
        "Your previous response was not valid JSON or did not match the required schema. "
        f"Return ONLY a JSON array of exactly {body.question_count} objects with keys "
        '"slot","question_type","difficulty","marks". Do not include any other text. '
        "Do not change the educational intent of the blueprint beyond what is necessary to satisfy the schema."
    )

    try:
        structured = await call_openrouter_structured_json(
            model=_QUIZ_MODEL,
            base_messages=base_messages,
            response_format=response_format,
            temperature_first=0.15,
            temperature_retry=0.0,
            retry_user_instruction=retry_instruction,
            validator=lambda parsed: validate_quiz_blueprint_items(parsed, expected_count=body.question_count),
            logger=logger,
            log_context=f"[quiz/blueprint][{session_id}]",
        )
        blueprint = structured.value
        for att in structured.attempts:
            usage = AIUsage(
                prompt_tokens=att.llm_result.usage.prompt_tokens,
                completion_tokens=att.llm_result.usage.completion_tokens,
                total_tokens=att.llm_result.usage.total_tokens,
                model=att.llm_result.usage.model,
                request_id=att.llm_result.usage.request_id,
            )
            await record_ai_usage(
                svc,
                user_id=user.id,
                source="quiz_builder",
                entity_type="quiz_session",
                entity_id=session_id,
                route=_blueprint_route,
                usage=usage,
                status="success" if att.status == "success" else "error",
                error_message=None if att.status == "success" else (att.error_message or att.status),
            )
    except StructuredJSONError as exc:
        logger.error("[quiz/blueprint][%s] structured output failed: %s", session_id, exc)
        for att in getattr(exc, "attempts", []) or []:
            usage = AIUsage(
                prompt_tokens=att.llm_result.usage.prompt_tokens,
                completion_tokens=att.llm_result.usage.completion_tokens,
                total_tokens=att.llm_result.usage.total_tokens,
                model=att.llm_result.usage.model,
                request_id=att.llm_result.usage.request_id,
            )
            await record_ai_usage(
                svc,
                user_id=user.id,
                source="quiz_builder",
                entity_type="quiz_session",
                entity_id=session_id,
                route=_blueprint_route,
                usage=usage,
                status="error",
                error_message=att.error_message or att.status,
            )
        raise HTTPException(status_code=502, detail="Could not create question plan. Please try again.") from exc
    except Exception as exc:
        logger.error("[quiz/blueprint][%s] LLM error: %s", session_id, exc)
        await record_ai_usage(
            svc,
            user_id=user.id,
            source="quiz_builder",
            entity_type="quiz_session",
            entity_id=session_id,
            route=_blueprint_route,
            status="error",
            error_message=str(exc),
        )
        raise HTTPException(status_code=502, detail="LLM service error") from exc

    def _save_blueprint() -> None:
        svc = get_service_client()
        svc.table("quiz_blueprints").delete().eq("session_id", session_id).execute()
        svc.table("quiz_blueprints").insert({
            "session_id": session_id,
            "blueprint_json": blueprint,
        }).execute()

    await asyncio.to_thread(_save_blueprint)
    return {"blueprint_json": blueprint}


# ── PATCH /v1/quiz/sessions/{session_id} ─────────────────────────────────────

class PatchSessionRequest(BaseModel):
    subject: str | None = None
    grade_level: str | None = None
    title: str | None = None
    preset: str | None = None


@router.patch("/sessions/{session_id}", summary="Update session metadata")
async def patch_session(
    session_id: str,
    body: PatchSessionRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    await asyncio.to_thread(_require_session, session_id, user.id)

    updates: dict[str, Any] = {}
    if body.subject is not None:
        updates["subject"] = body.subject or None
    if body.grade_level is not None:
        updates["grade_level"] = body.grade_level or None
    if body.title is not None:
        updates["title"] = body.title or None
    if body.preset is not None:
        if body.preset not in VALID_PRESETS:
            raise HTTPException(status_code=400, detail=f"Invalid preset: {body.preset!r}")
        updates["preset"] = body.preset

    if updates:
        def _update() -> None:
            svc = get_service_client()
            svc.table("quiz_builder_sessions").update(updates).eq("id", session_id).execute()
        await asyncio.to_thread(_update)

    return {"ok": True}


# ── PUT /v1/quiz/sessions/{session_id}/blueprint ──────────────────────────────

class SaveBlueprintRequest(BaseModel):
    blueprint_json: list[dict]
    difficulty_definitions: dict[str, str] | None = None


@router.put("/sessions/{session_id}/blueprint", summary="Save blueprint directly (no LLM)")
async def save_blueprint_direct(
    session_id: str,
    body: SaveBlueprintRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    await asyncio.to_thread(_require_session, session_id, user.id)

    def _save() -> None:
        svc = get_service_client()
        row: dict[str, Any] = {
            "session_id": session_id,
            "blueprint_json": body.blueprint_json,
        }
        if body.difficulty_definitions is not None:
            row["difficulty_definitions"] = body.difficulty_definitions
        svc.table("quiz_blueprints").delete().eq("session_id", session_id).execute()
        svc.table("quiz_blueprints").insert(row).execute()

    await asyncio.to_thread(_save)
    return {"ok": True}


# ── POST /v1/quiz/sessions/{session_id}/generate ──────────────────────────────

@router.post(
    "/sessions/{session_id}/generate",
    status_code=status.HTTP_202_ACCEPTED,
    summary="Enqueue quiz question generation",
)
async def start_generation(
    session_id: str,
    request: Request,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    session = await asyncio.to_thread(_require_session, session_id, user.id)

    def _now_iso() -> str:
        return datetime.now(timezone.utc).isoformat()

    def _parse_ts(ts: str | None) -> datetime | None:
        if not ts:
            return None
        try:
            return datetime.fromisoformat(ts.replace("Z", "+00:00"))
        except Exception:
            return None

    # Keep aligned with frontend stuck detection (NewSessionWizard) to avoid UI showing "stuck"
    # before the backend considers the job stale and eligible for retry/reset.
    GENERATION_STALE_SECONDS = 6 * 60  # must be < worker job_timeout (10 min)

    def _is_stale(sess: dict) -> bool:
        last = _parse_ts(sess.get("updated_at")) or _parse_ts(sess.get("generation_started_at"))
        if not last:
            return False
        age_s = (datetime.now(timezone.utc) - last).total_seconds()
        return age_s >= GENERATION_STALE_SECONDS

    in_progress = (session.get("generation_status") in {"queued", "generating"}) or (session.get("status") == "generating")
    stale = in_progress and _is_stale(session)
    if in_progress and not stale:
        logger.info(
            "[quiz/generate] rejected duplicate session=%s gen_status=%s gen_stage=%s job_id=%s",
            session_id,
            session.get("generation_status"),
            session.get("generation_stage"),
            session.get("generation_job_id"),
        )
        return {"enqueued": False, "already_running": True, "session_id": session_id, "job_id": session.get("generation_job_id")}

    redis = request.app.state.redis
    old_job_id = session.get("generation_job_id")
    if stale and old_job_id:
        try:
            from arq.jobs import Job

            aborted = await Job(job_id=old_job_id, redis=redis).abort(timeout=0.5)
            logger.warning("[quiz/generate] stale generation abort attempted session=%s old_job_id=%s aborted=%s", session_id, old_job_id, aborted)
        except Exception as exc:
            logger.warning("[quiz/generate] stale generation abort failed session=%s old_job_id=%s exc=%s", session_id, old_job_id, exc)

    now_iso = _now_iso()
    attempt = int(session.get("generation_attempt_count") or 0) + 1
    job_id = f"quiz_{session_id}_{attempt}"

    def _set_queued() -> None:
        svc = get_service_client()
        svc.table("quiz_builder_sessions").update({
            "status": "generating",
            "error_message": None,
            "generation_status": "queued",
            "generation_stage": "queued",
            "generation_job_id": job_id,
            "generation_attempt_count": attempt,
            "generation_started_at": now_iso,
            "generation_completed_at": None,
            "generation_failed_at": None,
            "generation_error_code": None,
            "generation_error_message": None,
            "updated_at": now_iso,
        }).eq("id", session_id).execute()

    await asyncio.to_thread(_set_queued)

    enqueued_job = await redis.enqueue_job("generate_quiz", session_id=session_id, _job_id=job_id)
    if enqueued_job is None:
        logger.info("[quiz/generate] enqueue returned None (duplicate job_id) session=%s job_id=%s", session_id, job_id)
        return {"enqueued": False, "already_running": True, "session_id": session_id, "job_id": job_id}

    logger.info("[quiz/generate] enqueued session=%s job_id=%s attempt=%d", session_id, job_id, attempt)
    return {"enqueued": True, "session_id": session_id, "job_id": job_id}


# ── Question helpers ───────────────────────────────────────────────────────────

def _require_question(session_id: str, question_id: str, teacher_id: str) -> dict:
    _require_session(session_id, teacher_id)
    svc = get_service_client()
    result = (
        svc.table("quiz_questions")
        .select("*")
        .eq("id", question_id)
        .eq("session_id", session_id)
        .maybe_single()
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Question not found")
    return result.data


# ── GET /v1/quiz/sessions/{session_id}/questions ──────────────────────────────

@router.get("/sessions/{session_id}/questions", summary="List questions for a session")
async def list_questions(
    session_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> list[dict]:
    await asyncio.to_thread(_require_session, session_id, user.id)

    def _query() -> list[dict]:
        svc = get_service_client()
        result = (
            svc.table("quiz_questions")
            .select("*")
            .eq("session_id", session_id)
            .order("slot")
            .execute()
        )
        return result.data or []

    return await asyncio.to_thread(_query)


# ── POST /v1/quiz/sessions/{session_id}/questions/reorder ─────────────────────
# Must be defined BEFORE /{question_id} routes to avoid path conflict

class ReorderRequest(BaseModel):
    ordered_ids: list[str]


@router.post("/sessions/{session_id}/questions/reorder", summary="Reorder questions by slot")
async def reorder_questions(
    session_id: str,
    body: ReorderRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    await asyncio.to_thread(_require_session, session_id, user.id)

    def _reorder() -> None:
        svc = get_service_client()
        for slot, question_id in enumerate(body.ordered_ids, start=1):
            svc.table("quiz_questions").update({"slot": slot}).eq("id", question_id).eq("session_id", session_id).execute()

    await asyncio.to_thread(_reorder)
    return {"ok": True}


# ── PATCH /v1/quiz/sessions/{session_id}/questions/{question_id} ─────────────

class PatchQuestionRequest(BaseModel):
    question_json: dict | None = None
    teacher_approved: bool | None = None


@router.patch("/sessions/{session_id}/questions/{question_id}", summary="Update a question")
async def update_question(
    session_id: str,
    question_id: str,
    body: PatchQuestionRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    existing = await asyncio.to_thread(_require_question, session_id, question_id, user.id)

    updates: dict[str, Any] = {}
    if body.question_json is not None:
        # Merge partial update into existing question_json so fields like slot,
        # question_type, bloom_level, rubric, acceptable_answers are preserved.
        merged = {**existing.get("question_json", {}), **body.question_json}
        _validate_mc_options(merged)
        updates["question_json"] = _normalize_mc_answer(merged)
    if body.teacher_approved is not None:
        updates["teacher_approved"] = body.teacher_approved

    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")

    def _update() -> dict:
        svc = get_service_client()
        result = (
            svc.table("quiz_questions")
            .update(updates)
            .eq("id", question_id)
            .execute()
        )
        return result.data[0] if result.data else {}

    return await asyncio.to_thread(_update)


# ── DELETE /v1/quiz/sessions/{session_id}/questions/{question_id} ─────────────

@router.delete(
    "/sessions/{session_id}/questions/{question_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete a question",
)
async def delete_question(
    session_id: str,
    question_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> Response:
    await asyncio.to_thread(_require_question, session_id, question_id, user.id)

    def _delete() -> None:
        svc = get_service_client()
        svc.table("quiz_questions").delete().eq("id", question_id).execute()

    await asyncio.to_thread(_delete)
    return Response(status_code=status.HTTP_204_NO_CONTENT)


# ── POST /v1/quiz/sessions/{session_id}/questions/{question_id}/regenerate ────

_REGENERATE_PROMPT = """\
You are an expert quiz creator. Regenerate a single quiz question.

Session context:
- Subject: {subject}
- Grade level: {grade_level}
- Mode: {mode}
- Preset: {preset}

Source material:
{source_text}

Difficulty definitions — use these when writing the question at the requested level:
- Easy: {easy_def}
- Medium: {medium_def}
- Challenge: {challenge_def}

LaTeX / KaTeX math & scientific notation (store LaTeX source text, not HTML) — REQUIRED whenever the source or question contains formula-like notation (even if Subject is "General").
Apply to ALL fields: question_text/options/correct_answer/acceptable_answers/explanation/rubric.
At minimum, question_text/options/correct_answer/explanation/rubric must follow these rules.
- Delimiters in the stored text: \\( ... \\) for inline, \\[ ... \\] for display.
- JSON escaping: you are returning JSON, so backslashes must be escaped. In your JSON output strings write \\\\( ... \\\\) and \\\\[ ... \\\\].
  Example JSON string value: "\\\\( \\\\frac{{d}}{{dx}}\\\\sin x = \\\\cos x \\\\)"
- Convert EVERY formula-like snippet (derivative, integral, limit, equation, symbol expression, unit expression, exponent/subscript, trig identity, scientific notation) into KaTeX-compatible LaTeX and wrap it in delimiters. Do NOT leave ASCII math.
- Forbidden ASCII math examples (do not output these): d/dx sin(x) = cos(x); sin^2(x) + cos^2(x) = 1
  Required LaTeX (stored text): \\( \\frac{{d}}{{dx}}\\sin x = \\cos x \\); \\( \\sin^2 x + \\cos^2 x = 1 \\)
- Use LaTeX syntax (KaTeX-friendly): fractions \\frac{{a}}{{b}}, roots \\sqrt{{x}}, subscripts x_{{n}}, superscripts x^{{n}}, Greek \\alpha \\beta \\gamma, multiplication \\times or \\cdot, units \\text{{m/s}}.
- Never output HTML tags (e.g., <span>...</span>) or KaTeX/MathJax-generated HTML; only LaTeX source inside delimiters.

Current question being replaced:
{current_question_json}

Keep slot={slot} fixed. For question_type, difficulty, and marks, use the defaults below; defer to the teacher's instruction if it explicitly requests a different value for any of these.
{clarifying_prompt_section}
Generate exactly ONE question. Teacher's instruction above overrides any of these defaults:
- slot: {slot} (never change)
- question_type: {question_type} (default; override if teacher requests)
- difficulty: {difficulty} (default; override if teacher requests)
- marks: {marks} (default; override if teacher requests)

Return a single JSON object (not an array) with this schema:
{{"slot": {slot}, \
"question_type": "multiple_choice"|"fill_blank"|"short_answer"|"long_answer", \
"difficulty": "easy"|"medium"|"challenge", \
"bloom_level": "remember"|"understand"|"apply"|"analyse"|"evaluate"|"create", \
"question_text": "<full question text>", \
"options": ["A. ...", "B. ...", "C. ...", "D. ..."] for MC else null, \
"correct_answer": "<answer>", "acceptable_answers": ["<alt answer>"] or null, "explanation": "<brief explanation>", \
"marks": <integer 1-5>, "rubric": "<rubric for long_answer>" or null, \
"answer_space": "multiple_choice"|"single_line"|"short_paragraph"|"long_paragraph", \
"source_reference": null, "image_slot": null}}

Return ONLY the raw JSON object. No markdown, no explanation."""


class RegenerateRequest(BaseModel):
    clarifying_prompt: str | None = None


_REGEN_RESPONSE_FORMAT: dict[str, Any] = {
    "type": "json_schema",
    "json_schema": {
        "name": "quiz_question",
        "strict": True,
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "slot": {"type": "integer", "minimum": 1},
                "question_type": {
                    "type": "string",
                    "enum": ["multiple_choice", "fill_blank", "short_answer", "long_answer"],
                },
                "difficulty": {"type": "string", "enum": ["easy", "medium", "challenge"]},
                "bloom_level": {
                    "type": "string",
                    "enum": ["remember", "understand", "apply", "analyse", "evaluate", "create"],
                },
                "question_text": {"type": "string", "minLength": 1},
                "options": {"type": ["array", "null"], "items": {"type": "string"}},
                "correct_answer": {"type": "string", "minLength": 1},
                "acceptable_answers": {"type": ["array", "null"], "items": {"type": "string"}},
                "explanation": {"type": "string", "minLength": 1},
                "marks": {"type": "integer", "minimum": 1, "maximum": 5},
                "rubric": {"type": ["string", "null"]},
                "answer_space": {
                    "type": "string",
                    "enum": ["multiple_choice", "single_line", "short_paragraph", "long_paragraph"],
                },
                "source_reference": {"type": ["string", "null"]},
                "image_slot": {"type": "null"},
            },
            "required": [
                "slot", "question_type", "difficulty", "bloom_level", "question_text",
                "options", "correct_answer", "acceptable_answers", "explanation",
                "marks", "rubric", "answer_space", "source_reference", "image_slot",
            ],
        },
    },
}


@router.post(
    "/sessions/{session_id}/questions/{question_id}/regenerate",
    summary="Regenerate a single question via LLM",
)
async def regenerate_question(
    session_id: str,
    question_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
    body: RegenerateRequest | None = None,
) -> dict:
    question = await asyncio.to_thread(_require_question, session_id, question_id, user.id)
    session = await asyncio.to_thread(_require_session, session_id, user.id)

    clarifying_prompt: str | None = None
    if body is not None and body.clarifying_prompt:
        cp = body.clarifying_prompt.strip()[:1000]
        clarifying_prompt = cp if cp else None

    def _get_source() -> dict | None:
        svc = get_service_client()
        result = (
            svc.table("quiz_source_materials")
            .select("extracted_text")
            .eq("session_id", session_id)
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        )
        return result.data[0] if result.data else None

    def _get_blueprint_defs() -> dict | None:
        svc = get_service_client()
        result = (
            svc.table("quiz_blueprints")
            .select("difficulty_definitions")
            .eq("session_id", session_id)
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        )
        return result.data[0] if result.data else None

    source = await asyncio.to_thread(_get_source)
    bp = await asyncio.to_thread(_get_blueprint_defs)
    source_text: str = (source or {}).get("extracted_text") or "(no source text provided)"
    raw_defs = (bp or {}).get("difficulty_definitions") or {}
    defs = {
        "easy": raw_defs.get("easy") or _DEFAULT_DIFFICULTY_DEFS["easy"],
        "medium": raw_defs.get("medium") or _DEFAULT_DIFFICULTY_DEFS["medium"],
        "challenge": raw_defs.get("challenge") or _DEFAULT_DIFFICULTY_DEFS["challenge"],
    }

    qj = question.get("question_json", {})
    current_question_json = json.dumps(qj, indent=2, ensure_ascii=False)
    clarifying_prompt_section = (
        f"Teacher's instruction: {clarifying_prompt}\n\n"
        if clarifying_prompt
        else ""
    )
    prompt = _REGENERATE_PROMPT.format(
        subject=session.get("subject") or "General",
        grade_level=session.get("grade_level") or "Secondary",
        mode=session.get("mode", "from_prompt"),
        preset=session.get("preset", "standard_worksheet"),
        source_text=source_text[:3000],
        easy_def=defs["easy"],
        medium_def=defs["medium"],
        challenge_def=defs["challenge"],
        question_type=qj.get("question_type", "short_answer"),
        difficulty=qj.get("difficulty", "medium"),
        marks=qj.get("marks", 2),
        slot=question.get("slot", 1),
        current_question_json=current_question_json,
        clarifying_prompt_section=clarifying_prompt_section,
    )

    svc_regen = get_service_client()
    _regen_route = f"POST /v1/quiz/sessions/{session_id}/questions/{question_id}/regenerate"
    base_messages = [
        {"role": "system", "content": "Return only valid JSON matching the required schema. No markdown fences, no prose."},
        {"role": "user", "content": prompt},
    ]
    retry_instruction = (
        "Your previous response failed JSON/schema validation. "
        "Fix only the JSON formatting/structure and field values needed for schema compliance; "
        "do not change the educational content beyond what is necessary. "
        "Return ONLY the corrected JSON object."
    )

    try:
        structured = await call_openrouter_structured_json(
            model=_QUIZ_MODEL,
            base_messages=base_messages,
            response_format=_REGEN_RESPONSE_FORMAT,
            temperature_first=0.3,
            temperature_retry=0.0,
            retry_user_instruction=retry_instruction,
            validator=lambda parsed: QuizQuestionSchema.model_validate(parsed).model_dump(),
            logger=logger,
            log_context=f"[quiz/regenerate][{session_id}/{question_id}]",
        )
        new_qj: dict[str, Any] = structured.value
        for att in structured.attempts:
            usage = AIUsage(
                prompt_tokens=att.llm_result.usage.prompt_tokens,
                completion_tokens=att.llm_result.usage.completion_tokens,
                total_tokens=att.llm_result.usage.total_tokens,
                model=att.llm_result.usage.model,
                request_id=att.llm_result.usage.request_id,
            )
            await record_ai_usage(
                svc_regen,
                user_id=user.id,
                source="quiz_builder",
                entity_type="quiz_question",
                entity_id=question_id,
                route=_regen_route,
                usage=usage,
                status="success" if att.status == "success" else "error",
                error_message=None if att.status == "success" else (att.error_message or att.status),
            )
    except StructuredJSONError as exc:
        logger.error("[quiz/regenerate][%s/%s] structured output failed: %s", session_id, question_id, exc)
        for att in getattr(exc, "attempts", []) or []:
            usage = AIUsage(
                prompt_tokens=att.llm_result.usage.prompt_tokens,
                completion_tokens=att.llm_result.usage.completion_tokens,
                total_tokens=att.llm_result.usage.total_tokens,
                model=att.llm_result.usage.model,
                request_id=att.llm_result.usage.request_id,
            )
            await record_ai_usage(
                svc_regen,
                user_id=user.id,
                source="quiz_builder",
                entity_type="quiz_question",
                entity_id=question_id,
                route=_regen_route,
                usage=usage,
                status="error",
                error_message=att.error_message or att.status,
            )
        raise HTTPException(status_code=502, detail="Could not regenerate question") from exc
    except Exception as exc:
        logger.error("[quiz/regenerate][%s/%s] error: %s", session_id, question_id, exc)
        await record_ai_usage(
            svc_regen,
            user_id=user.id,
            source="quiz_builder",
            entity_type="quiz_question",
            entity_id=question_id,
            route=_regen_route,
            status="error",
            error_message=str(exc),
        )
        raise HTTPException(status_code=502, detail="Could not regenerate question") from exc

    def _update() -> dict:
        svc = get_service_client()
        result = (
            svc.table("quiz_questions")
            .update({
                "question_json": new_qj,
                "validation_status": "pending",
                "validation_warnings": [],
                "teacher_approved": False,
            })
            .eq("id", question_id)
            .execute()
        )
        return result.data[0] if result.data else {}

    return await asyncio.to_thread(_update)


# ── POST /v1/quiz/sessions/{session_id}/questions/{question_id}/clone ─────────

@router.post(
    "/sessions/{session_id}/questions/{question_id}/clone",
    status_code=status.HTTP_201_CREATED,
    summary="Clone a question",
)
async def clone_question(
    session_id: str,
    question_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    question = await asyncio.to_thread(_require_question, session_id, question_id, user.id)

    def _clone() -> dict:
        svc = get_service_client()
        max_result = (
            svc.table("quiz_questions")
            .select("slot")
            .eq("session_id", session_id)
            .order("slot", desc=True)
            .limit(1)
            .execute()
        )
        max_slot = max_result.data[0]["slot"] if max_result.data else question["slot"]
        insert_result = svc.table("quiz_questions").insert({
            "session_id": session_id,
            "slot": max_slot + 1,
            "question_json": question["question_json"],
            "validation_status": question["validation_status"],
            "validation_warnings": question.get("validation_warnings") or [],
            "teacher_approved": False,
        }).execute()
        return insert_result.data[0] if insert_result.data else {}

    return await asyncio.to_thread(_clone)


# ── POST /v1/quiz/sessions/{session_id}/questions/{question_id}/image ─────────

_QUIZ_IMAGES_BUCKET = "quiz-images"


@router.post(
    "/sessions/{session_id}/questions/{question_id}/image",
    summary="Upload an image for a question",
)
async def upload_question_image(
    session_id: str,
    question_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
    file: Annotated[UploadFile, File()],
    position: Annotated[str, Form()] = "above",
    caption: Annotated[str, Form()] = "",
    size: Annotated[str, Form()] = "medium",
) -> dict:
    await asyncio.to_thread(_require_question, session_id, question_id, user.id)

    file_bytes = await file.read()
    filename = file.filename or "image.jpg"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "jpg"
    storage_path = f"{session_id}/{question_id}/{uuid.uuid4()}.{ext}"
    content_type = file.content_type or "image/jpeg"

    def _upload() -> str:
        svc = get_service_client()
        svc.storage.from_(_QUIZ_IMAGES_BUCKET).upload(
            path=storage_path,
            file=file_bytes,
            file_options={"content-type": content_type},
        )
        return f"{settings.supabase_url}/storage/v1/object/public/{_QUIZ_IMAGES_BUCKET}/{storage_path}"

    image_url = await asyncio.to_thread(_upload)
    return {"image_url": image_url, "storage_path": storage_path}


# ── DELETE /v1/quiz/sessions/{session_id}/questions/{question_id}/image ───────

@router.delete(
    "/sessions/{session_id}/questions/{question_id}/image",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Remove an image from a question",
)
async def remove_question_image(
    session_id: str,
    question_id: str,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
    storage_path: Annotated[str | None, Query()] = None,
) -> Response:
    # storage_path query param: used for orphan cleanup when a file was uploaded but
    # the follow-up PATCH failed. Must be scoped to this session/question to prevent
    # deletion of other questions' assets.
    if storage_path is not None:
        expected_prefix = f"{session_id}/{question_id}/"
        if not storage_path.startswith(expected_prefix):
            raise HTTPException(status_code=403, detail="storage_path out of scope")
        await asyncio.to_thread(_require_question, session_id, question_id, user.id)
    else:
        question = await asyncio.to_thread(_require_question, session_id, question_id, user.id)
        qj = question.get("question_json", {})
        image_slot = qj.get("image_slot")
        storage_path = (image_slot or {}).get("storage_path") if image_slot else None

    if storage_path:
        def _remove() -> None:
            svc = get_service_client()
            svc.storage.from_(_QUIZ_IMAGES_BUCKET).remove([storage_path])

        try:
            await asyncio.to_thread(_remove)
        except Exception as exc:
            logger.warning("[quiz/image/remove] storage remove failed: %s", exc)

    return Response(status_code=status.HTTP_204_NO_CONTENT)


# ── Source text extraction helpers ─────────────────────────────────────────────

def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        import io as _io
        from pypdf import PdfReader
        reader = PdfReader(_io.BytesIO(file_bytes))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p.strip() for p in parts if p.strip())
    except Exception as exc:
        logger.warning("[quiz/source] PDF extraction failed: %s", exc)
        return ""


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        import io as _io
        from docx import Document as _DocxDocument
        doc = _DocxDocument(_io.BytesIO(file_bytes))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(parts)
    except Exception as exc:
        logger.warning("[quiz/source] DOCX extraction failed: %s", exc)
        return ""


# ── HTML / Markdown rendering helpers ─────────────────────────────────────────

_ANSWER_LINE_COUNTS: dict[str, int] = {"none": 0, "small": 3, "medium": 6, "large": 12}

# KaTeX CDN assets included in the PDF/preview HTML so math renders before Playwright prints.
# delimiters match what the LLM is instructed to emit: \( inline \) and \[ display \].
# window.__PDF_READY is set after renderMathInElement completes; render_html_to_pdf_bytes waits for it.
_KATEX_HEAD = (
    '<link rel="stylesheet"'
    ' href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css"'
    ' crossorigin="anonymous">'
    '<script defer'
    ' src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"'
    ' crossorigin="anonymous"></script>'
    '<script defer'
    ' src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"'
    ' crossorigin="anonymous"></script>'
    '<script>'
    'document.addEventListener("DOMContentLoaded",function(){'
    'if(typeof renderMathInElement==="function"){'
    'renderMathInElement(document.body,{'
    'delimiters:[{left:"\\\\[",right:"\\\\]",display:true},'
    '{left:"\\\\(",right:"\\\\)",display:false}],'
    'throwOnError:false});'
    '}else{'
    '(window.__PDF_WARNINGS=window.__PDF_WARNINGS||[]).push("KaTeX auto-render not loaded");'
    '}'
    'window.__PDF_READY=true;'
    '});'
    '</script>'
)


def _h(v: object) -> str:
    """HTML-escape any value for safe inline interpolation."""
    return _html.escape(str(v or ""))


def _is_quiz_image_url(url: str, session_id: str, question_id: str = "") -> bool:
    """Return True only if the URL is within this question's quiz-images path.
    The uploaded path structure is {session_id}/{question_id}/{uuid}.{ext}, so
    scoping to both IDs prevents a question from embedding an image that was
    uploaded for a different question in the same quiz."""
    if not url or not session_id or not question_id:
        return False
    prefix = f"{settings.supabase_url}/storage/v1/object/public/quiz-images/{session_id}/{question_id}/"
    return url.startswith(prefix)


def _get_accessible_class_ids(svc: Any, user_id: str, role: str) -> set[str]:
    """Return class IDs the caller is allowed to access, based on their role."""
    if role in ("teacher", "tutor"):
        rows = svc.table("classes").select("id").eq("teacher_id", user_id).execute().data or []
        return {r["id"] for r in rows}
    if role == "co_teacher":
        rows = (
            svc.table("class_co_teachers")
            .select("class_id")
            .eq("user_id", user_id)
            .not_.is_("accepted_at", "null")
            .execute().data or []
        )
        return {r["class_id"] for r in rows}
    if role == "school_admin":
        user_row = (
            svc.table("users").select("workspace_id").eq("id", user_id).maybe_single().execute().data
        )
        workspace_id = (user_row or {}).get("workspace_id")
        if not workspace_id:
            return set()
        rows = svc.table("classes").select("id").eq("workspace_id", workspace_id).execute().data or []
        return {r["id"] for r in rows}
    return set()


_QUIZ_CSS = """
* { box-sizing: border-box; margin: 0; padding: 0; }
body { font-family: 'Times New Roman', serif; font-size: 11pt; line-height: 1.5; color: #000; padding: 18mm 20mm; }
.header { border-bottom: 2px solid #000; padding-bottom: 10px; margin-bottom: 18px; }
.quiz-title { font-size: 16pt; font-weight: bold; }
.quiz-subject { font-size: 10pt; color: #444; margin-top: 4px; }
.fields { display: flex; gap: 36px; margin-top: 10px; }
.field { font-size: 10pt; display: flex; align-items: center; gap: 6px; }
.field-line { border-bottom: 1px solid #000; width: 160px; display: inline-block; }
.question { margin-bottom: 20px; page-break-inside: avoid; }
.question-header { display: flex; align-items: flex-start; gap: 6px; }
.question-num { font-weight: bold; min-width: 22px; flex-shrink: 0; }
.question-text { flex: 1; }
.marks { font-size: 9pt; color: #666; white-space: nowrap; margin-left: 8px; flex-shrink: 0; }
.options { margin: 8px 0 0 28px; display: flex; flex-direction: column; gap: 3px; }
.option { padding: 1px 0; }
.option.correct { font-weight: bold; background: #ffffc0; padding: 1px 4px; border-radius: 2px; }
.answer-space { margin: 8px 0 0 28px; }
.answer-line { border-bottom: 1px solid #bbb; height: 1px; margin-bottom: 18px; }
.teacher-answer { margin: 5px 0 0 28px; font-size: 9pt; color: #005bb5; font-style: italic;
                  background: #e8f4fd; border-left: 3px solid #005bb5; padding: 3px 8px; border-radius: 2px; }
.question-image { margin: 8px 0 6px 28px; }
.question-image img { display: block; height: auto; border: 1px solid #ddd; border-radius: 2px; }
.image-caption { font-size: 9pt; color: #666; margin-top: 3px; font-style: italic; }
@media print { @page { size: A4; margin: 18mm 20mm; } body { padding: 0; } }
"""

_QUIZ_CSS_OVERRIDES: dict[str, str] = {
    "exam_style": (
        "body { font-family: Arial, sans-serif; } "
        ".question { border-top: 1px solid #ccc; padding-top: 12px; } "
        ".marks { font-weight: bold; }"
    ),
    "worksheet_answer_lines": ".answer-line { margin-bottom: 22px; }",
    "compact_quiz": ".question { margin-bottom: 12px; } .answer-line { margin-bottom: 10px; }",
    "reading_comprehension": ".question-text { font-style: italic; }",
}


def _render_question_html(
    number: int,
    qj: dict,
    answer_space_override: str,
    show_marks: bool,
    is_teacher: bool,
    session_id: str = "",
    question_id: str = "",
) -> str:
    q_type = qj.get("question_type", "short_answer")
    text = _h(qj.get("question_text") or "")
    marks = qj.get("marks", 1)
    options = qj.get("options") or []
    correct = _h(qj.get("correct_answer") or "")
    explanation = _h(qj.get("explanation") or "")
    answer_space_field = qj.get("answer_space", "short_paragraph")

    raw_slot = qj.get("image_slot")
    image_slot: dict = raw_slot if isinstance(raw_slot, dict) else {}
    img_url = image_slot.get("url", "")
    img_pos = image_slot.get("position", "above")
    img_caption = _h(image_slot.get("caption") or "")
    img_width = {"small": "200px", "medium": "350px", "large": "500px"}.get(
        image_slot.get("size", "medium"), "350px"
    )

    space_map = {
        "multiple_choice": "none",
        "single_line": "small",
        "short_paragraph": "medium",
        "long_paragraph": "large",
    }
    if answer_space_override in ("none", "small", "large"):
        space_key = answer_space_override
    else:
        space_key = space_map.get(answer_space_field, "medium")

    n_lines = _ANSWER_LINE_COUNTS.get(space_key, 6)
    marks_html = f'<span class="marks">[{marks} mark{"s" if marks != 1 else ""}]</span>' if show_marks else ""

    def _img_html() -> str:
        if not img_url:
            return ""
        caption_html = f'<div class="image-caption">{img_caption}</div>' if img_caption else ""
        return (
            f'<div class="question-image">'
            f'<img src="{_h(img_url)}" style="max-width:{img_width};height:auto" alt="{img_caption}">'
            f"{caption_html}</div>"
        )

    html = '<div class="question"><div class="question-header">'
    html += f'<span class="question-num">{number}.</span>'
    html += '<div class="question-body">'
    if img_url and _is_quiz_image_url(img_url, session_id, question_id) and img_pos == "above":
        html += _img_html()
    html += f'<span class="question-text">{text}</span>'
    if img_url and _is_quiz_image_url(img_url, session_id, question_id) and img_pos == "below":
        html += _img_html()
    html += f"{marks_html}</div></div>"

    if q_type == "multiple_choice" and options:
        html += '<div class="options">'
        for opt in options:
            is_correct = is_teacher and bool(correct) and opt.startswith(correct[0])
            cls = "option correct" if is_correct else "option"
            html += f'<div class="{cls}">{_h(opt)}</div>'
        html += "</div>"
    else:
        html += '<div class="answer-space">' + '<div class="answer-line"></div>' * n_lines + "</div>"

    if is_teacher:
        html += f'<div class="teacher-answer">Answer: {correct}</div>'
        if explanation:
            html += f'<div class="teacher-answer">Explanation: {explanation}</div>'
        rubric = _h(qj.get("rubric") or "")
        if rubric:
            html += f'<div class="teacher-answer">Rubric: {rubric}</div>'

    html += "</div>"
    return html


def _generate_quiz_html(
    session: dict,
    questions: list[dict],
    template_id: str,
    config: dict,
    version: str = "student",
) -> str:
    title = session.get("title") or "Quiz"
    subject = session.get("subject") or ""
    show_marks = config.get("show_marks", True)
    show_name = config.get("show_name_field", True)
    show_date = config.get("show_date_field", True)
    answer_space_override = config.get("answer_space", "medium")
    is_teacher = version == "teacher"

    css = _QUIZ_CSS + _QUIZ_CSS_OVERRIDES.get(template_id, "")

    fields_parts: list[str] = []
    if show_name:
        fields_parts.append('<div class="field">Name:&nbsp;<span class="field-line">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></div>')
    if show_date:
        fields_parts.append('<div class="field">Date:&nbsp;<span class="field-line">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span></div>')

    header = f'<div class="header"><div class="quiz-title">{_h(title)}</div>'
    if subject:
        header += f'<div class="quiz-subject">{_h(subject)}</div>'
    if fields_parts:
        header += '<div class="fields">' + "".join(fields_parts) + "</div>"
    header += "</div>"

    sid = session.get("id") or ""
    sorted_qs = sorted(questions, key=lambda q: q.get("slot", 99))
    qs_html = "".join(
        _render_question_html(
            i + 1, q.get("question_json", {}), answer_space_override, show_marks, is_teacher,
            session_id=sid, question_id=q.get("id") or "",
        )
        for i, q in enumerate(sorted_qs)
    )

    return (
        '<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">'
        f"<style>{css}</style>"
        f"{_KATEX_HEAD}"
        f"</head><body>"
        f"{header}<div class='questions'>{qs_html}</div></body></html>"
    )


def _generate_quiz_markdown(
    session: dict,
    questions: list[dict],
    version: str = "student",
) -> str:
    title = session.get("title") or "Quiz"
    subject = session.get("subject") or ""
    is_teacher = version == "teacher"

    lines: list[str] = [f"# {title}"]
    if subject:
        lines.append(f"**Subject:** {subject}")
    lines.append("")

    sid = session.get("id") or ""
    sorted_qs = sorted(questions, key=lambda q: q.get("slot", 99))
    for i, q in enumerate(sorted_qs):
        qj = q.get("question_json", {})
        q_num = i + 1
        q_type = qj.get("question_type", "short_answer")
        marks = qj.get("marks", 1)
        text = qj.get("question_text", "")
        options = qj.get("options") or []
        correct = qj.get("correct_answer", "")
        explanation = qj.get("explanation", "")

        raw_slot = qj.get("image_slot")
        image_slot: dict = raw_slot if isinstance(raw_slot, dict) else {}
        img_url = image_slot.get("url", "")
        img_caption = image_slot.get("caption", "")
        img_pos = image_slot.get("position", "above")
        qid = q.get("id") or ""
        img_md = (
            f"![{img_caption}]({img_url})"
            if img_url and _is_quiz_image_url(img_url, sid, qid)
            else ""
        )

        lines.append(f"## {q_num}. {text}  *[{marks} mark{'s' if marks != 1 else ''}]*")
        lines.append("")
        if img_md and img_pos == "above":
            lines.append(img_md)
            lines.append("")
        if q_type == "multiple_choice" and options:
            for opt in options:
                lines.append(f"- {opt}")
            lines.append("")
        else:
            lines.append("_" * 60)
            lines.append("")
        if img_md and img_pos != "above":
            lines.append(img_md)
            lines.append("")
        if is_teacher:
            lines.append(f"> **Answer:** {correct}")
            if explanation:
                lines.append(f"> *{explanation}*")
            lines.append("")

    return "\n".join(lines)


# ── Render + Export DB helpers ─────────────────────────────────────────────────

_QUIZ_EXPORTS_BUCKET = "quiz-exports"


def _list_questions_sync(session_id: str, *, approved_only: bool = False) -> list[dict]:
    svc = get_service_client()
    q = svc.table("quiz_questions").select("*").eq("session_id", session_id)
    if approved_only:
        q = q.eq("teacher_approved", True)
    return q.order("slot").execute().data or []


def _store_render_sync(session_id: str, template_id: str, config: dict) -> str:
    svc = get_service_client()
    result = svc.table("quiz_renders").insert({
        "session_id": session_id,
        "template_id": template_id,
        "config": config,
    }).execute()
    return result.data[0]["id"] if result.data else str(uuid.uuid4())


def _get_render_sync(render_id: str) -> dict | None:
    svc = get_service_client()
    return svc.table("quiz_renders").select("*").eq("id", render_id).maybe_single().execute().data


# ── POST /v1/quiz/sessions/{session_id}/render ────────────────────────────────

class RenderConfigModel(BaseModel):
    answer_space: str = "medium"
    show_marks: bool = True
    show_name_field: bool = True
    show_date_field: bool = True
    page_size: str = "A4"
    include_instructions: bool = True
    title: str = ""


class RenderRequest(BaseModel):
    template_id: str
    render_config: RenderConfigModel = RenderConfigModel()


@router.post("/sessions/{session_id}/render", summary="Render quiz to HTML preview")
async def render_quiz(
    session_id: str,
    body: RenderRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    session = await asyncio.to_thread(_require_session, session_id, user.id)
    questions = await asyncio.to_thread(_list_questions_sync, session_id, approved_only=True)
    if not questions:
        raise HTTPException(status_code=400, detail="No approved questions to render")
    config = body.render_config.model_dump()
    html_preview = _generate_quiz_html(session, questions, body.template_id, config, "student")
    render_id = await asyncio.to_thread(_store_render_sync, session_id, body.template_id, config)
    return {"render_id": render_id, "html_preview": html_preview}


# ── POST /v1/quiz/sessions/{session_id}/export/pdf ───────────────────────────

class ExportBinaryRequest(BaseModel):
    version: str = "student"
    render_id: str | None = None


@router.post("/sessions/{session_id}/export/pdf", summary="Export quiz as PDF via Playwright")
async def export_pdf(
    session_id: str,
    body: ExportBinaryRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    session = await asyncio.to_thread(_require_session, session_id, user.id)

    def _get_data() -> tuple[list[dict], dict | None]:
        qs = _list_questions_sync(session_id, approved_only=True)
        render = _get_render_sync(body.render_id) if body.render_id else None
        return qs, render

    questions, render_row = await asyncio.to_thread(_get_data)
    if not questions:
        raise HTTPException(status_code=400, detail="No approved questions to export")
    template_id = (render_row or {}).get("template_id", "exam_style")
    config = (render_row or {}).get("config", {})
    html = _generate_quiz_html(session, questions, template_id, config, body.version)

    pdf_bytes, pdf_warnings = await render_html_to_pdf_bytes(html)
    if pdf_warnings:
        logger.warning("[quiz/export_pdf][%s] PDF warnings: %s", session_id, pdf_warnings)

    file_id = body.render_id or str(uuid.uuid4())
    storage_path = f"{session_id}/{file_id}_{body.version}.pdf"

    def _upload_pdf() -> dict:
        svc = get_service_client()
        try:
            svc.storage.from_(_QUIZ_EXPORTS_BUCKET).upload(
                path=storage_path,
                file=pdf_bytes,
                file_options={"content-type": "application/pdf"},
            )
        except Exception:
            svc.storage.from_(_QUIZ_EXPORTS_BUCKET).update(
                path=storage_path,
                file=pdf_bytes,
                file_options={"content-type": "application/pdf"},
            )
        return svc.storage.from_(_QUIZ_EXPORTS_BUCKET).create_signed_url(storage_path, 3600)

    signed = await asyncio.to_thread(_upload_pdf)
    url = signed.get("signedURL") or signed.get("signedUrl") or signed.get("signed_url") or ""
    return {"signed_url": url}


# ── POST /v1/quiz/sessions/{session_id}/export/docx ──────────────────────────

@router.post("/sessions/{session_id}/export/docx", summary="Export quiz as DOCX")
async def export_docx(
    session_id: str,
    body: ExportBinaryRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    import io as _io
    import httpx
    from docx import Document as _DocxDoc
    from docx.shared import Inches

    session = await asyncio.to_thread(_require_session, session_id, user.id)

    def _get_docx_data() -> tuple[list[dict], dict | None]:
        qs = _list_questions_sync(session_id, approved_only=True)
        render = _get_render_sync(body.render_id) if body.render_id else None
        return qs, render

    questions, render_row = await asyncio.to_thread(_get_docx_data)
    if not questions:
        raise HTTPException(status_code=400, detail="No approved questions to export")
    config = (render_row or {}).get("config", {})

    show_marks = config.get("show_marks", True)
    show_name = config.get("show_name_field", True)
    show_date = config.get("show_date_field", True)
    answer_space_override = config.get("answer_space", "medium")
    is_teacher = body.version == "teacher"

    # Pre-fetch question images (public URLs from quiz-images bucket)
    sorted_qs = sorted(questions, key=lambda q: q.get("slot", 99))
    img_cache: dict[str, bytes] = {}
    async with httpx.AsyncClient(timeout=10.0) as client:
        for q in sorted_qs:
            raw_slot = q.get("question_json", {}).get("image_slot")
            slot: dict = raw_slot if isinstance(raw_slot, dict) else {}
            url = slot.get("url", "")
            q_session_id = q.get("session_id", "")
            q_question_id = q.get("id", "")
            if url and _is_quiz_image_url(url, q_session_id, q_question_id) and url not in img_cache:
                try:
                    resp = await client.get(url)
                    if resp.status_code == 200:
                        img_cache[url] = resp.content
                except Exception as exc:
                    logger.warning("[quiz/export/docx] image fetch failed url=%s: %s", url, exc)

    doc = _DocxDoc()
    title = session.get("title") or "Quiz"
    doc.add_heading(title, 0)
    if session.get("subject"):
        doc.add_paragraph(f"Subject: {session['subject']}")
    if show_name:
        doc.add_paragraph("Name: " + "_" * 35)
    if show_date:
        doc.add_paragraph("Date: " + "_" * 20)
    doc.add_paragraph()

    for i, q in enumerate(sorted_qs):
        qj = q.get("question_json", {})
        q_num = i + 1
        q_type = qj.get("question_type", "short_answer")
        marks = qj.get("marks", 1)
        text = qj.get("question_text", "")
        options = qj.get("options") or []
        correct = qj.get("correct_answer", "")
        explanation = qj.get("explanation", "")
        raw_image_slot = qj.get("image_slot")
        image_slot: dict = raw_image_slot if isinstance(raw_image_slot, dict) else {}
        img_url = image_slot.get("url", "")
        img_pos = image_slot.get("position", "above")

        def _add_image(doc: Any, url: str) -> None:
            img_bytes = img_cache.get(url)
            if img_bytes:
                try:
                    doc.add_picture(_io.BytesIO(img_bytes), width=Inches(3.5))
                except Exception as exc:
                    logger.warning("[quiz/export/docx] add_picture failed: %s", exc)

        if img_url and img_pos == "above":
            _add_image(doc, img_url)
            if image_slot.get("caption"):
                doc.add_paragraph(image_slot["caption"])

        marks_str = f" [{marks} mark{'s' if marks != 1 else ''}]" if show_marks else ""
        p = doc.add_paragraph()
        run = p.add_run(f"{q_num}. {text}{marks_str}")
        run.bold = True

        if img_url and img_pos == "below":
            _add_image(doc, img_url)
            if image_slot.get("caption"):
                doc.add_paragraph(image_slot["caption"])

        if q_type == "multiple_choice" and options:
            for opt in options:
                doc.add_paragraph(f"    {opt}")
        else:
            n_lines = _ANSWER_LINE_COUNTS.get(answer_space_override, 6)
            for _ in range(n_lines):
                doc.add_paragraph("_" * 70)

        if is_teacher:
            ans_p = doc.add_paragraph(f"[Answer: {correct}]")
            ans_p.runs[0].italic = True
            if explanation:
                exp_p = doc.add_paragraph(f"[{explanation}]")
                exp_p.runs[0].italic = True
        doc.add_paragraph()

    buf = _io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    file_id = body.render_id or str(uuid.uuid4())
    storage_path = f"{session_id}/{file_id}_{body.version}.docx"
    docx_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def _upload_docx() -> dict:
        svc = get_service_client()
        try:
            svc.storage.from_(_QUIZ_EXPORTS_BUCKET).upload(
                path=storage_path,
                file=docx_bytes,
                file_options={"content-type": docx_ct},
            )
        except Exception:
            svc.storage.from_(_QUIZ_EXPORTS_BUCKET).update(
                path=storage_path,
                file=docx_bytes,
                file_options={"content-type": docx_ct},
            )
        return svc.storage.from_(_QUIZ_EXPORTS_BUCKET).create_signed_url(storage_path, 3600)

    signed = await asyncio.to_thread(_upload_docx)
    url = signed.get("signedURL") or signed.get("signedUrl") or signed.get("signed_url") or ""
    return {"signed_url": url}


# ── POST /v1/quiz/sessions/{session_id}/export/markdown ──────────────────────

class MarkdownExportRequest(BaseModel):
    version: str = "student"
    render_id: str | None = None


@router.post("/sessions/{session_id}/export/markdown", summary="Export quiz as Markdown text")
async def export_markdown(
    session_id: str,
    body: MarkdownExportRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    session = await asyncio.to_thread(_require_session, session_id, user.id)
    questions = await asyncio.to_thread(_list_questions_sync, session_id, approved_only=True)
    if not questions:
        raise HTTPException(status_code=400, detail="No approved questions to export")
    return {"markdown": _generate_quiz_markdown(session, questions, body.version)}


# ── Personalization routes ─────────────────────────────────────────────────────

_PERSONALIZATION_PROMPT = """\
You are an educational analyst. Generate an anonymized student learning profile for quiz personalization.

Performance data:
{performance_data}

Signals to emphasize: {signals}

Write 2–3 concise paragraphs describing:
- Weak areas and common error patterns
- Strengths and areas of confidence
- Recent learning trend (improving, plateauing, or declining)

Completely anonymized — no names or IDs. Use "this student" / "the learner".
Return only the profile text, no JSON or headers."""


class PersonalizationRequest(BaseModel):
    student_id: str
    class_id: str | None = None
    selected_signals: list[str] = []


@router.post(
    "/sessions/{session_id}/personalization/summary",
    summary="Generate anonymized student performance summary",
)
async def generate_personalization_summary(
    session_id: str,
    body: PersonalizationRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    await asyncio.to_thread(_require_session, session_id, user.id)

    class _AuthError(Exception):
        pass

    def _fetch_performance() -> list[dict]:
        svc = get_service_client()

        # Resolve accessible classes for this caller; varies by role so that
        # co_teachers (class_co_teachers) and school_admins (workspace) work
        # correctly without leaking data across role boundaries.
        accessible_ids = _get_accessible_class_ids(svc, user.id, user.role)

        if body.class_id:
            if body.class_id not in accessible_ids:
                raise _AuthError()
            scope_class_ids = [body.class_id]
        else:
            if not accessible_ids:
                return []
            scope_class_ids = list(accessible_ids)

        asgn_ids = [
            r["id"]
            for r in (
                svc.table("assignments").select("id")
                .in_("class_id", scope_class_ids)
                .execute().data or []
            )
        ]
        if not asgn_ids:
            return []

        sub_ids = [
            r["id"]
            for r in (
                svc.table("submissions").select("id")
                .eq("student_id", body.student_id)
                .in_("assignment_id", asgn_ids)
                .order("submitted_at", desc=True)
                .limit(100)
                .execute().data or []
            )
        ]
        if not sub_ids:
            return []

        return (
            svc.table("grading_results")
            .select("marks_awarded, marks_available")
            .in_("submission_id", sub_ids[:100])
            .execute()
            .data or []
        )

    try:
        rows = await asyncio.to_thread(_fetch_performance)
    except _AuthError:
        raise HTTPException(status_code=403, detail="Class not found")

    if not rows:
        return {
            "anonymized_summary": (
                "No grading data found for this student yet. "
                "Consider using a prompt-based quiz or adding source material."
            )
        }

    valid = [r for r in rows if r.get("marks_available")]
    total = len(valid)
    avg_pct = (
        sum((r["marks_awarded"] or 0) / r["marks_available"] * 100 for r in valid) / total
        if total else 0
    )
    below_50 = sum(
        1 for r in valid if (r["marks_awarded"] or 0) / r["marks_available"] < 0.5
    )

    performance_data = (
        f"Questions graded: {total}. "
        f"Average score: {avg_pct:.1f}%. "
        f"Questions answered below 50%: {below_50}."
    )
    signals = body.selected_signals or ["weak_topics", "common_mistakes", "recent_mastery_trend"]
    prompt = _PERSONALIZATION_PROMPT.format(
        performance_data=performance_data,
        signals=", ".join(signals),
    )

    svc_pers = get_service_client()
    _pers_route = f"POST /v1/quiz/sessions/{session_id}/personalization"
    pers_ai_usage: AIUsage | None = None
    try:
        llm_result = await call_openrouter(
            model=_QUIZ_MODEL,
            temperature=0.5,
            messages=[{"role": "user", "content": prompt}],
        )
        pers_ai_usage = AIUsage(
            prompt_tokens=llm_result.usage.prompt_tokens,
            completion_tokens=llm_result.usage.completion_tokens,
            total_tokens=llm_result.usage.total_tokens,
            model=llm_result.usage.model,
            request_id=llm_result.usage.request_id,
        )
        summary = llm_result.content.strip()
        await record_ai_usage(
            svc_pers,
            user_id=user.id,
            source="quiz_builder",
            entity_type="quiz_session",
            entity_id=session_id,
            route=_pers_route,
            usage=pers_ai_usage,
        )
    except Exception as exc:
        logger.error("[quiz/personalization][%s] LLM error: %s", session_id, exc)
        await record_ai_usage(
            svc_pers,
            user_id=user.id,
            source="quiz_builder",
            entity_type="quiz_session",
            entity_id=session_id,
            route=_pers_route,
            status="error",
            error_message=str(exc),
        )
        summary = (
            f"This student has answered {total} questions with an average score of {avg_pct:.1f}%. "
            "Consider focusing on topics where performance falls below 50%."
        )

    return {"anonymized_summary": summary}


class SaveSummaryRequest(BaseModel):
    teacher_edited_summary: str


@router.patch(
    "/sessions/{session_id}/personalization/summary",
    summary="Save teacher-edited personalization summary to session",
)
async def save_personalization_summary(
    session_id: str,
    body: SaveSummaryRequest,
    user: Annotated[AuthUser, Depends(require_roles(*_TEACHER_ROLES))],
) -> dict:
    await asyncio.to_thread(_require_session, session_id, user.id)

    def _save() -> None:
        svc = get_service_client()
        svc.table("quiz_builder_sessions").update({
            "personalization_data": {"teacher_summary": body.teacher_edited_summary},
        }).eq("id", session_id).execute()

    await asyncio.to_thread(_save)
    return {"ok": True}
